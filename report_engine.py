"""
report_engine.py
================
Pipeline DSP + generazione .docx estratta da prova_lift_single.py.
Nessuna dipendenza da tkinter o matplotlib GUI backend.
"""

import io
import math
import numpy as np
import matplotlib
matplotlib.use("Agg")   # backend non-interattivo, obbligatorio su server
import matplotlib.figure as mpl_fig

# ---------------------------------------------------------------------------
# Costanti
# ---------------------------------------------------------------------------
KDL             = 1.22e-3
N_TRAX          = 8
FS              = 1250
DC_OFFSET       = 2.5
COEFF_CORR      = 20.833333

PEAK_FACTOR     = 2.5
PEAK_WIN_SEC    = 0.5
PEAK_MAX_IN_WIN = 2
ENV_WIN_DETECT  = 500
ENV_STEP_DETECT = 125
REF_WIN_SEC     = 5.0

RMS_WIN_DETECT  = 250
RMS_STEP_DETECT = 125
MIN_ACTIVE_SEC  = 3.0
REF_SEC         = 1.0
FACTOR_THRESH   = 2.0
ACTIVE_RATIO    = 0.10
INITIAL_TRIM_SEC= 0.5
EXCLUDE_END_SEC = 5.0

Y_LIM_MIN = -0.15
Y_LIM_MAX =  0.15
COLOR_VALID    = "#d62728"
COLOR_EXCLUDED = "#1f77b4"

AMC_DRAWER   = "Ing. Matteo Ressia"
AMC_APPROVAL = "Ing. Bruno Vusini"
AMC_FOOTER   = (
    "AMC Instruments s.r.l.\n"
    "Sede legale Corso Matteotti, 36 - 10121 Torino (Italy)\n"
    "Sede operativa Via Pietro Nenni 79/E - 10036 Settimo Torinese (TO) (Italy)\n"
    "P.I. 09612820010 - REA 1066557 - Capitale Sociale € 14.244,87\n"
    "Tel: +39 011 0378820 - Fax: +39 011 19835584 - info@aemmeci.com - www.aemmeci.com"
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _sig_display(raw_ch):
    return (raw_ch - DC_OFFSET) / COEFF_CORR


# ---------------------------------------------------------------------------
# Lettura .bin
# ---------------------------------------------------------------------------

def read_bin(data: bytes) -> np.ndarray:
    """Legge bytes grezzi del file .bin -> array (N, N_TRAX)."""
    n_uint16 = len(data) // 2
    flat = np.frombuffer(data[:n_uint16 * 2], dtype="<u2").astype(np.float64) * KDL
    n_samples = len(flat) // N_TRAX
    result = np.zeros((n_samples, N_TRAX))
    for k in range(N_TRAX):
        result[:, k] = flat[k: n_samples * N_TRAX: N_TRAX]
    return result


# ---------------------------------------------------------------------------
# Parsing .txt
# ---------------------------------------------------------------------------

FIRMWARE_LABELS = {
    "Firmware1.3.1": {
        "labels": [
            "Test_count:", "Date and time:", "Matricola impianto(*):",
            "Tipologia Impianto(*):", "Velocita' funi [m/s]:", "Numero funi(*):",
            "Diametro funi(*):", "Numero trefoli:", "Tipologia anima:",
            "Citta'(*):", "Provincia(*):", "Piazza/Via/Corso(*):",
            "Civico(*):", "Interno/Scala:", "CAP:",
            "Anno installaz. funi:", "Numero fermate:", "Ultimo responso:",
        ],
        "map": {
            "testcount": 0, "matricola": 2, "tipimp": 3, "velocita": 4,
            "nfuni": 5, "diametro": 6, "trefoli": 7, "anima": 8,
            "citta": 9, "provincia": 10, "indirizzo": 11, "civico": 12,
            "internoscala": 13, "cap": 14, "anno": 15, "nfermate": 16,
            "responso": 17,
        },
    },
    "Firmware1.3": {
        "labels": [
            "Test_count:", "Date and time:", "Matricola impianto(*):",
            "Tipologia Impianto(*):", "Velocita' funi [m/s]:", "Numero funi(*):",
            "Diametro funi(*):", "Numero trefoli:", "Tipologia anima:",
            "Citta'(*):", "Provincia(*):", "Piazza/Via/Corso(*):",
            "Civico(*):", "Interno/Scala:", "CAP:",
            "Anno installazione:", "Numero piani:", "Ultimo responso:",
        ],
        "map": {
            "testcount": 0, "matricola": 2, "tipimp": 3, "velocita": 4,
            "nfuni": 5, "diametro": 6, "trefoli": 7, "anima": 8,
            "citta": 9, "provincia": 10, "indirizzo": 11, "civico": 12,
            "internoscala": 13, "cap": 14, "anno": 15, "nfermate": 16,
            "responso": 17,
        },
    },
}

_FW_134_LABELS = [
    "DLH sn:", "Test_count:", "Date and time:", "Matricola impianto(*):",
    "Tipologia Impianto(*):", "Velocita' funi [m/s]:", "Numero funi(*):",
    "Diametro funi(*):", "Numero trefoli:", "Tipologia anima:",
    "Produttore fune:", "Citta'(*):", "Provincia(*):", "Piazza/Via/Corso(*):",
    "Civico(*):", "Interno/Scala:", "CAP:", "Anno installaz. funi:",
    "Numero fermate:", "Esito ultimo responso:", "Note:", "Barcode:",
]
_FW_134_MAP = {
    "DLHsn": 0, "testcount": 1, "matricola": 3, "tipimp": 4, "velocita": 5,
    "nfuni": 6, "diametro": 7, "trefoli": 8, "anima": 9,
    "citta": 11, "provincia": 12, "indirizzo": 13, "civico": 14,
    "internoscala": 15, "cap": 16, "anno": 17, "nfermate": 18,
    "responso": 19, "Note": 20, "Barcode": 21,
}
_FW_11X_LABELS = [
    "DLH sn:", "Test_count:", "Date and time:", "Rapporto di intervento(*):",
    "Tipologia Impianto(*):", "Velocita' funi [m/s]:", "Numero funi(*):",
    "Diametro funi(*):", "Numero trefoli:", "Tipologia anima:",
    "Produttore fune:", "Anno installaz. funi:", "Numero fermate:",
    "Esito ultimo responso:",
]
_FW_11X_MAP = {
    "testcount": 1, "matricola": 3, "tipimp": 4, "velocita": 5,
    "nfuni": 6, "diametro": 7, "trefoli": 8, "anima": 9,
    "anno": 11, "nfermate": 12, "responso": 13,
}
_FW_OLD = {
    "Dati impianto ed acquisizione:": [
        "Matricola impianto(*):", "Tipologia Impianto(*):", "Numero funi(*):",
        "Diametro funi(*):", "Citta'(*):", "Provincia(*):",
        "Piazza/Via/Corso(*):", "Civico(*):", "Interno:", "Scala:",
        "CAP:", "Anno installazione:", "Numero piani:",
        "Ente notificato:", "Ultimo responso:",
    ],
    "Firmware1.1": [
        "Matricola impianto(*):", "Tipologia Impianto(*):",
        "Velocita' funi [m/s]:", "Numero funi(*):", "Diametro funi(*):",
        "numero trefoli:", "Tipologia anima:", "Citta'(*):", "Provincia(*):",
        "Piazza/Via/Corso(*):", "Civico(*):", "Interno/Scala:", "CAP:",
        "Anno installazione:", "Numero piani:", "Ultimo responso:",
    ],
}


def _parse_labels(lines, labels, n=None):
    if n is None:
        n = len(labels)
    answ = [""] * len(labels)
    idx = 0
    for i, lbl in enumerate(labels[:n]):
        while idx < len(lines):
            if lines[idx].strip().lower() == lbl.lower():
                answ[i] = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
                idx += 2
                break
            idx += 1
    return answ


def _g(answ, i, fallback="N.R."):
    v = answ[i] if i < len(answ) else ""
    return v if v else fallback


def parse_txt(txt_bytes: bytes) -> dict:
    lines = txt_bytes.decode("utf-8", errors="replace").splitlines()
    fw = lines[0].strip() if lines else "N.R."
    r = {
        "firmware": fw,
        "matricola": "N.R.", "tipimp": "N.R.", "velocita": "N.R.",
        "nfuni": "N.R.", "diametro": "N.R.", "trefoli": "N.R.",
        "anima": "N.R.", "citta": "N.R.", "provincia": "N.R.",
        "indirizzo": "N.R.", "civico": "N.R.", "internoscala": "N.R.",
        "cap": "N.R.", "anno": "N.R.", "nfermate": "N.R.",
        "responso": "N.R.", "testcount": "N.R.", "DLHsn": "N.R.",
        "Note": "N.R.", "Barcode": "N.R.",
    }
    if fw in FIRMWARE_LABELS:
        cfg = FIRMWARE_LABELS[fw]
        a = _parse_labels(lines, cfg["labels"])
        for k, i in cfg["map"].items():
            r[k] = _g(a, i)
    elif fw in ("Firmware1.3.2", "Firmware1.3.3", "Firmware1.3.4") or fw.startswith("Firmware1.8"):
        n = len(_FW_134_LABELS) if fw.startswith("Firmware1.8") else len(_FW_134_LABELS) - 2
        a = _parse_labels(lines, _FW_134_LABELS, n)
        for k, i in _FW_134_MAP.items():
            r[k] = _g(a, i)
    elif fw.startswith("Firmware1.1."):
        a = _parse_labels(lines, _FW_11X_LABELS)
        for k, i in _FW_11X_MAP.items():
            r[k] = _g(a, i)
    elif fw in _FW_OLD:
        a = _parse_labels(lines, _FW_OLD[fw])
        if fw == "Dati impianto ed acquisizione:":
            r["matricola"] = _g(a, 0); r["tipimp"] = _g(a, 1)
            r["nfuni"] = _g(a, 2); r["diametro"] = _g(a, 3)
            r["citta"] = _g(a, 4); r["provincia"] = _g(a, 5)
            r["indirizzo"] = _g(a, 6); r["civico"] = _g(a, 7)
            r["internoscala"] = _g(a, 8, "") + _g(a, 9, "")
            r["cap"] = _g(a, 10); r["anno"] = _g(a, 11)
            r["nfermate"] = _g(a, 12); r["responso"] = _g(a, 14)
            r["velocita"] = "0.7"
        else:
            r["matricola"] = _g(a, 0); r["tipimp"] = _g(a, 1)
            r["velocita"] = _g(a, 2); r["nfuni"] = _g(a, 3)
            r["diametro"] = _g(a, 4); r["trefoli"] = _g(a, 5)
            r["anima"] = _g(a, 6); r["citta"] = _g(a, 7)
            r["provincia"] = _g(a, 8); r["indirizzo"] = _g(a, 9)
            r["civico"] = _g(a, 10); r["internoscala"] = _g(a, 11)
            r["cap"] = _g(a, 12); r["anno"] = _g(a, 13)
            r["nfermate"] = _g(a, 14); r["responso"] = _g(a, 15)
    return r


# ---------------------------------------------------------------------------
# DSP
# ---------------------------------------------------------------------------

def detect_active_channels(lf, n_funi_expected, n_ch=6):
    n_ch = min(n_ch, lf.shape[1])
    rms_ac = np.array([
        math.sqrt(float(np.mean((lf[:, ch] - np.mean(lf[:, ch])) ** 2)))
        if lf.shape[0] > 0 else 0.0
        for ch in range(n_ch)
    ])
    if n_funi_expected <= 0 or n_funi_expected >= n_ch:
        rms_max = np.max(rms_ac) if np.max(rms_ac) > 0 else 1.0
        return rms_ac >= ACTIVE_RATIO * rms_max
    ranked = np.argsort(rms_ac)[::-1]
    active = np.zeros(n_ch, dtype=bool)
    active[ranked[:n_funi_expected]] = True
    return active


def detect_shared_zone(lf, active_mask, n_ch=6):
    import statistics
    n_samp    = lf.shape[0]
    fallback  = (0, n_samp - 1)
    min_f     = max(1, int(math.ceil(MIN_ACTIVE_SEC * FS / RMS_STEP_DETECT)))
    max_gap_f = int(5.0 * FS / RMS_STEP_DETECT)
    ref_frames = max(1, int(REF_SEC * FS / RMS_STEP_DETECT))
    envs = {}
    for ch in range(min(n_ch, lf.shape[1])):
        if ch < len(active_mask) and active_mask[ch]:
            sig_ac = np.abs(lf[:, ch] - float(np.mean(lf[:, ch])))
            starts = np.arange(0, n_samp - RMS_WIN_DETECT + 1, RMS_STEP_DETECT)
            if len(starts) == 0:
                continue
            mad     = np.array([float(np.mean(sig_ac[s: s + RMS_WIN_DETECT])) for s in starts])
            centers = starts + RMS_WIN_DETECT // 2
            envs[ch] = (centers, mad, float(np.mean(mad[:ref_frames])))
    if not envs:
        return [fallback] * n_ch
    quietest_ch = min(envs, key=lambda c: envs[c][2])
    ref_thr = FACTOR_THRESH * envs[quietest_ch][2]
    zone_per_ch = {}
    for ch, (centers, mad, _) in envs.items():
        above = mad >= ref_thr
        runs, in_run, rs = [], False, 0
        for i, v in enumerate(above):
            if v and not in_run:
                in_run, rs = True, i
            elif not v and in_run:
                in_run = False
                runs.append((rs, i - 1))
        if in_run:
            runs.append((rs, len(above) - 1))
        valid = [(s, e) for s, e in runs if (e - s + 1) >= min_f]
        if not valid:
            zone_per_ch[ch] = fallback
            continue
        merged, cur = [], list(valid[0])
        for s, e in valid[1:]:
            if s - cur[1] <= max_gap_f:
                cur[1] = e
            else:
                merged.append(tuple(cur)); cur = [s, e]
        merged.append(tuple(cur))
        best = max(merged, key=lambda r: r[1] - r[0])
        ss = max(0,          int(centers[best[0]]) - RMS_WIN_DETECT // 2)
        es = min(n_samp - 1, int(centers[best[1]]) + RMS_WIN_DETECT // 2)
        zone_per_ch[ch] = (ss, es)
    all_starts = [z[0] for z in zone_per_ch.values()]
    all_stops  = [z[1] for z in zone_per_ch.values()]
    shared_s, shared_e = max(all_starts), min(all_stops)
    if shared_e <= shared_s:
        shared_s = int(statistics.median(all_starts))
        shared_e = int(statistics.median(all_stops))
    return [(shared_s, shared_e) if (ch < len(active_mask) and active_mask[ch])
            else fallback for ch in range(n_ch)]


def _mad_envelope(signal):
    sig_ac = np.abs(signal - float(np.mean(signal)))
    n = len(sig_ac)
    starts = np.arange(0, n - ENV_WIN_DETECT + 1, ENV_STEP_DETECT)
    if len(starts) == 0:
        return np.full(n, float(np.mean(sig_ac)))
    mad     = np.array([float(np.mean(sig_ac[s: s + ENV_WIN_DETECT])) for s in starts])
    centers = starts + ENV_WIN_DETECT // 2
    return np.interp(np.arange(n), centers, mad)


def _p99_envelope(signal):
    sig_ac = np.abs(signal - float(np.mean(signal)))
    n = len(sig_ac)
    starts = np.arange(0, n - ENV_WIN_DETECT + 1, ENV_STEP_DETECT)
    if len(starts) == 0:
        return np.full(n, float(np.percentile(sig_ac, 99)))
    p99     = np.array([float(np.percentile(sig_ac[s: s + ENV_WIN_DETECT], 99)) for s in starts])
    centers = starts + ENV_WIN_DETECT // 2
    return np.interp(np.arange(n), centers, p99)


def _base_level_channel(signal, ref_win_sec=REF_WIN_SEC):
    env = _p99_envelope(signal)
    n_samp = len(env)
    ref_samp     = max(1, int(ref_win_sec * FS))
    exclude_samp = int(EXCLUDE_END_SEC * FS)
    search_end   = max(ref_samp, n_samp - exclude_samp)
    if ref_samp >= search_end:
        return max(float(np.mean(env[:search_end])), 1e-6)
    best_sum = float('inf')
    best_mean = float(np.mean(env[:search_end]))
    for i in range(0, search_end - ref_samp + 1, ENV_STEP_DETECT):
        s = float(np.sum(env[i: i + ref_samp]))
        if s < best_sum:
            best_sum  = s
            best_mean = float(np.mean(env[i: i + ref_samp]))
    return max(best_mean, 1e-6)


def analyze_peaks(lf, active_mask=None, valid_zones=None,
                  peak_factor=PEAK_FACTOR, ref_win_sec=REF_WIN_SEC):
    n_ch   = min(6, lf.shape[1])
    n_samp = lf.shape[0]
    win_s  = int(PEAK_WIN_SEC * FS)
    step_s = max(1, int(0.05 * FS))
    per_canale, esiti_attivi, base_levels = [], [], []
    for ch in range(n_ch):
        is_active = (active_mask is None or
                     (ch < len(active_mask) and bool(active_mask[ch])))
        if not is_active:
            per_canale.append({'esito': 'inattivo', 'n_picchi': 0,
                               'n_concentrazioni': 0, 'peak_indices': np.array([])})
            continue
        s_v, e_v = (valid_zones[ch] if (valid_zones and ch < len(valid_zones))
                    else (0, n_samp - 1))
        seg = lf[s_v: e_v + 1, ch]
        if len(seg) < ENV_WIN_DETECT:
            per_canale.append({'esito': 'OK', 'n_picchi': 0,
                               'n_concentrazioni': 0, 'peak_indices': np.array([]),
                               'valid_start': s_v, 'base_level': 0.0})
            esiti_attivi.append('OK'); base_levels.append(0.0)
            continue
        base_level = _base_level_channel(seg, ref_win_sec)
        threshold  = peak_factor * base_level
        base_levels.append(base_level)
        seg_ac = np.abs(seg - float(np.mean(seg)))
        above  = seg_ac > threshold
        MERGE_GAP = int(0.08 * FS)
        runs_raw, in_p, ps = [], False, 0
        for i, v in enumerate(above):
            if v and not in_p:
                in_p, ps = True, i
            elif not v and in_p:
                in_p = False; runs_raw.append([ps, i - 1])
        if in_p:
            runs_raw.append([ps, len(above) - 1])
        merged_runs = []
        if runs_raw:
            cur = list(runs_raw[0])
            for r in runs_raw[1:]:
                if r[0] - cur[1] <= MERGE_GAP:
                    cur[1] = r[1]
                else:
                    merged_runs.append(tuple(cur)); cur = list(r)
            merged_runs.append(tuple(cur))
        peaks = np.array([s + int(np.argmax(seg_ac[s: e + 1]))
                          for s, e in merged_runs], dtype=int)
        n_pk, n_conc = len(peaks), 0
        if n_pk > PEAK_MAX_IN_WIN:
            for ws in range(0, max(1, len(seg) - win_s + 1), step_s):
                if int(np.sum((peaks >= ws) & (peaks < ws + win_s))) > PEAK_MAX_IN_WIN:
                    n_conc += 1
        esito_ch = ('KO_concentrazione' if n_conc > 0
                    else 'KO_singolo' if n_pk > 0 else 'OK')
        esiti_attivi.append(esito_ch)
        per_canale.append({'esito': esito_ch, 'n_picchi': n_pk,
                           'n_concentrazioni': n_conc,
                           'peak_indices': peaks + s_v,
                           'valid_start': s_v, 'base_level': base_level})
    priority = {'KO_concentrazione': 2, 'KO_singolo': 1, 'OK': 0}
    esito_globale = (max(esiti_attivi, key=lambda e: priority.get(e, -1))
                     if esiti_attivi else 'OK')
    return {'esito_globale': esito_globale, 'per_canale': per_canale,
            'base_level': float(np.mean(base_levels)) if base_levels else 0.0}


def trim_initial_from_peaks(peak_result, active_mask, zones,
                            trim_sec=INITIAL_TRIM_SEC, n_ch=6):
    trim_samples = int(trim_sec * FS)
    if trim_samples <= 0:
        return zones, False
    per_ch = peak_result.get("per_canale", [])
    needs_trim = False
    for ch in range(min(n_ch, len(zones))):
        if ch >= len(active_mask) or not bool(active_mask[ch]):
            continue
        if ch >= len(per_ch):
            continue
        s_v, _ = zones[ch]
        peaks = per_ch[ch].get("peak_indices", np.array([]))
        if len(peaks) and np.any((peaks >= s_v) & (peaks < s_v + trim_samples)):
            needs_trim = True; break
    if not needs_trim:
        return zones, False
    new_zones = []
    for ch in range(len(zones)):
        if ch < len(active_mask) and bool(active_mask[ch]):
            s_v, e_v = zones[ch]
            new_zones.append((min(s_v + trim_samples, e_v), e_v))
        else:
            new_zones.append(zones[ch])
    return new_zones, True


# ---------------------------------------------------------------------------
# Render PNG canale
# ---------------------------------------------------------------------------

def render_channel_png(lf, t, ch, s_v, e_v, esito_ch,
                       width_in=14, height_in=3, dpi=110) -> io.BytesIO:
    sig = _sig_display(lf[:, ch])
    fig = mpl_fig.Figure(figsize=(width_in, height_in), tight_layout=True)
    ax  = fig.add_subplot(111)
    ax.set_facecolor("white")
    if s_v > 0:
        ax.plot(t[:s_v], sig[:s_v], color=COLOR_EXCLUDED, linewidth=0.7, alpha=0.6)
    ax.plot(t[s_v: e_v+1], sig[s_v: e_v+1], color=COLOR_VALID, linewidth=0.9)
    if e_v + 1 < len(t):
        ax.plot(t[e_v+1:], sig[e_v+1:], color=COLOR_EXCLUDED, linewidth=0.7, alpha=0.6)
    for x_m in (t[s_v], t[min(e_v, len(t)-1)]):
        ax.axvline(x=x_m, color="#555555", linewidth=0.9, linestyle=":", alpha=0.7)
    ax.axhline(y=0, color="#888888", linewidth=0.5, alpha=0.5)
    col_t = "#d62728" if "KO" in esito_ch else "#2ca02c"
    ax.set_title(f"Ch {ch+1}  [{esito_ch}]", fontsize=10, color=col_t, fontweight="bold")
    ax.set_ylim(Y_LIM_MIN, Y_LIM_MAX)
    ax.set_xlim(0, t[-1] if len(t) > 0 else 1)
    ax.set_xlabel("Time [s]", fontsize=8)
    ax.set_ylabel("LF [V]", fontsize=8)
    ax.tick_params(labelsize=7)
    ax.grid(True, alpha=0.25, linewidth=0.4)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    fig.clf(); del fig
    return buf


# ---------------------------------------------------------------------------
# Generazione .docx  ->  restituisce io.BytesIO
# ---------------------------------------------------------------------------

def generate_report(meta: dict, lf: np.ndarray, t: np.ndarray,
                    active_mask: np.ndarray, valid_zones: list,
                    peak_result: dict, selected_channels: list,
                    esito_globale_manual: str,
                    test_date_str: str, doc_number: str) -> io.BytesIO:

    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21);   section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5);  section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5);  section.right_margin  = Cm(2.5)

    def _par(text="", align=WD_ALIGN_PARAGRAPH.CENTER,
             bold=False, size=11, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
        return p

    def _heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def _tc(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size)

    addr    = meta.get("indirizzo", "N.R.")
    civico  = meta.get("civico", "")
    citta   = meta.get("citta", "N.R.")
    prov    = meta.get("provincia", "")
    matr    = meta.get("matricola", "N.R.")
    tipimp  = meta.get("tipimp", "N.R.")
    n_funi  = meta.get("nfuni", "N.R.")

    addr_full  = f"{addr}, {civico}".strip(", ") if addr not in ("N.R.", "") else "N.R."
    citta_full = f"{citta} ({prov})".strip() if citta not in ("N.R.", "") else "N.R."

    # --- Frontespizio ---
    _par("SERVICE REPORT", bold=True, size=22, space_before=40, space_after=30)
    _par("VERIFICA MAGNETO-INDUTTIVA DELLE FUNI TRAENTI SU IMPIANTO ASCENSORISTICO\nPRESSO STABILE SITUATO IN:",
         bold=False, size=16, space_before=10, space_after=10)
    _par(addr_full.upper(), bold=True, size=18, space_before=4, space_after=4)
    _par(citta_full.upper(), bold=True, size=18, space_before=0, space_after=20)
    _par(f"IMPIANTO  {matr}", bold=True, size=14, space_before=10, space_after=6)
    _par(f"TIPO: {tipimp}", bold=False, size=12, space_before=0, space_after=30)
    _par("AMC INSTRUMENTS", bold=True, size=20, space_before=10, space_after=10)

    doc.add_paragraph()
    tbl_c = doc.add_table(rows=2, cols=4)
    tbl_c.style = 'Table Grid'
    tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_c.cell(0, 0).merge(tbl_c.cell(1, 0))
    _tc(tbl_c.cell(0, 0),
        "All proprietary rights reserved by AMC instruments Srl. "
        "This document shall not be reproduced or utilized by third parties "
        "without an AMC written consent.", size=7)
    _tc(tbl_c.cell(0, 1), "Doc. N°", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tc(tbl_c.cell(0, 2), doc_number, size=9)
    tbl_c.cell(0, 3).merge(tbl_c.cell(1, 3))
    _tc(tbl_c.cell(0, 3), "Pag. tot.: —", size=9)
    _tc(tbl_c.cell(1, 1), "DRAWER / APPROVAL", bold=True, size=8)
    _tc(tbl_c.cell(1, 2),
        f"TEST DATE:  {test_date_str}\n"
        f"DRAWER:     {AMC_DRAWER}\n"
        f"APPROVAL:   {AMC_APPROVAL}", size=8)
    _par(AMC_FOOTER, size=7, space_before=20, space_after=0)

    # --- Pag 2: Introduzione + Criteri + Tabella ---
    doc.add_page_break()
    _heading("1  INTRODUZIONE")
    p = doc.add_paragraph(
        "La presente relazione riporta i risultati ottenuti dalle prove condotte direttamente "
        "sull'impianto. I test eseguiti hanno la funzione di rilevare lo stato delle funi traenti "
        "dell'ascensore presente nello stabile visitato. Il test è consistito nell'analisi "
        "magneto-induttiva delle funi metalliche dove le eventuali anomalie rilevate indicano "
        "la presenza di fili rotti interni e/o esterni alla fune metallica.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _heading("2  CRITERI DI ANALISI")
    p = doc.add_paragraph(
        f"Il numero di funi traenti dell'ascensore sotto analisi è pari a {n_funi}. "
        "La strumentazione adottata per tale analisi è denominata AMC LIFT LC06H-TK. "
        "Essa consente un test magneto-induttivo non contemporaneo di tutte le funi traenti. "
        "La strumentazione fornisce all'operatore un grafico (traccia) significativo dello stato "
        "della fune; in particolare l'indicazione dei fili rotti è evidenziata dalla presenza di "
        "picchi della traccia. Si considerano difetti significativi quelli associati ad un "
        "determinato rapporto segnale/rumore SNR. Nel seguito sono riportate le tracce per "
        "ciascuna delle funi verificate.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _heading("3  DATI CARATTERISTICI DELL'IMPIANTO IN ANALISI")
    rows_data = [
        ("NUMERO DI IMPIANTO",               matr),
        ("INDIRIZZO",                         addr_full),
        ("CITTÀ",                             citta),
        ("PROVINCIA",                         prov),
        ("CAP",                               meta.get("cap", "N.R.")),
        ("DIAMETRO FUNI TRAENTI (mm)",        meta.get("diametro", "N.R.")),
        ("NUMERO DI FUNI TRAENTI",            n_funi),
        ("NUMERO DI TREFOLI",                 meta.get("trefoli", "N.R.")),
        ("ANNO DI INSTALLAZIONE DELLE FUNI",  meta.get("anno", "N.R.")),
        ("TIPOLOGIA IMPIANTO",                tipimp),
        ("RESPONSO ULTIMA VERIFICA SEMESTRALE", meta.get("responso", "N.R.")),
        ("NOTE",                              meta.get("Note", "")),
    ]
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        cl, cr = tbl.cell(i, 0), tbl.cell(i, 1)
        cl.width = Cm(9); cr.width = Cm(7)
        _tc(cl, label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _tc(cr, value, size=9)
    _par(AMC_FOOTER, size=7, space_before=30, space_after=0)

    # --- Pag 3: Grafici ---
    doc.add_page_break()
    _heading("4  Risultati dei test")
    per_ch_list = peak_result.get("per_canale", [])
    for rope_num, ch in enumerate(selected_channels, start=1):
        if ch >= lf.shape[1]:
            continue
        s_v, e_v = valid_zones[ch] if ch < len(valid_zones) else (0, len(t) - 1)
        esito_ch = (per_ch_list[ch].get("esito", "OK") if ch < len(per_ch_list) else "OK")
        buf = render_channel_png(lf, t, ch, s_v, e_v, esito_ch)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(buf, width=Cm(16))
        cap = doc.add_paragraph(f"Figura {rope_num}: Segnale LF per fune {rope_num}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cap.runs[0].font.size = Pt(9); cap.runs[0].italic = True
    _par(AMC_FOOTER, size=7, space_before=20, space_after=0)

    # --- Pag 4: Conclusioni ---
    doc.add_page_break()
    _heading("5  Conclusioni")
    addr_conc = f"{addr_full.upper()} - {matr} - {citta_full.upper()}"
    p = doc.add_paragraph(
        f"Dall'analisi dei diagrammi riportati al paragrafo precedente è possibile rilevare "
        f"lo stato di integrità delle funi in verifica presso lo stabile situato in {addr_conc}. "
        f"L'analisi dei segnali, eseguita da un tecnico abilitato secondo normativa ISO EN 9712, "
        f"ha evidenziato quanto segue:")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    eg = peak_result.get("esito_globale", "OK")
    if esito_globale_manual == "PROVA RIUSCITA":
        checked = 1
    elif esito_globale_manual == "SOSTITUIRE":
        checked = 3
    else:
        checked = {"OK": 1, "KO_singolo": 2, "KO_concentrazione": 3}.get(eg, 2)

    CHECK, EMPTY = "☑", "☐"
    options = [
        (1, f"Le funi nel tratto analizzato risultano integre: non si rilevano difetti significativi."),
        (2, f"Le funi nel tratto analizzato presentano difetti ma in misura inferiore a quanto previsto "
            f"dalla normativa ISO 4344. Le funi possono pertanto essere mantenute in opera ma si consiglia "
            f"di eseguire un test mediante metodo magnetoinduttivo con cadenza annuale."),
        (3, f"Le funi nel tratto analizzato presentano difetti in misura superiore a quanto previsto "
            f"dalla normativa ISO 4344. Le funi debbono essere sostituite."),
    ]
    for num, text in options:
        p_opt = doc.add_paragraph(f"{CHECK if checked == num else EMPTY}  {text}")
        p_opt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_opt.paragraph_format.space_after = Pt(8)
        p_opt.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p_disc = doc.add_paragraph(
        "La valutazione complessiva dello stato di integrità delle funi, finalizzata alla loro "
        "eventuale sostituzione, va correlata con la conoscenza dettagliata dell'impianto, con "
        "la sua storia e con quella delle funi. Per tale motivo essa è demandata al personale "
        "di manutenzione che può fare le proprie valutazioni anche sulla base del presente "
        "esame magneto-induttivo.")
    p_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.add_run("AMC Instruments srl\n").bold = True
    p_sign.add_run("Dr. Ing. Bruno Vusini\n\n")
    r = p_sign.add_run("Iscrizione all'Ordine degli ingegneri di Asti n. A552\n"
                        "III Liv. MIT secondo ISO 9712\n"
                        "Cert. RINA N. 11MI5PO01")
    r.font.size = Pt(9)
    _par(AMC_FOOTER, size=7, space_before=20, space_after=0)

    # Salva in memoria
    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out


# ---------------------------------------------------------------------------
# Pipeline completa: bytes bin + bytes txt -> BytesIO docx
# ---------------------------------------------------------------------------

def run_pipeline(bin_bytes: bytes, txt_bytes: bytes,
                 selected_channels: list | None,
                 esito_manual: str,
                 test_date_str: str,
                 doc_number: str,
                 peak_factor: float = PEAK_FACTOR,
                 ref_win_sec: float = REF_WIN_SEC) -> io.BytesIO:

    lf   = read_bin(bin_bytes)
    meta = parse_txt(txt_bytes)
    N    = lf.shape[0]
    t    = np.linspace(1, N, N) / FS

    try:
        n_funi_exp = int(meta.get("nfuni", "0"))
    except (ValueError, TypeError):
        n_funi_exp = 0

    active = detect_active_channels(lf, n_funi_exp, n_ch=6)
    zones  = detect_shared_zone(lf, active, n_ch=6)
    pr     = analyze_peaks(lf, active, zones, peak_factor=peak_factor, ref_win_sec=ref_win_sec)
    zones, trimmed = trim_initial_from_peaks(pr, active, zones)
    if trimmed:
        pr = analyze_peaks(lf, active, zones, peak_factor=peak_factor, ref_win_sec=ref_win_sec)

    if selected_channels is None:
        selected_channels = [ch for ch in range(6) if ch < len(active) and active[ch]]

    return generate_report(
        meta=meta, lf=lf, t=t,
        active_mask=active, valid_zones=zones, peak_result=pr,
        selected_channels=selected_channels,
        esito_globale_manual=esito_manual,
        test_date_str=test_date_str,
        doc_number=doc_number,
    )
